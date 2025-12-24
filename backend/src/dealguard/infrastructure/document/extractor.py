"""Document text extraction for PDF and DOCX files."""

import hashlib
import io
from dataclasses import dataclass, field
from typing import cast

import fitz  # PyMuPDF
from docx import Document

from dealguard.shared.exceptions import UnsupportedFileTypeError, ValidationError
from dealguard.shared.logging import get_logger

logger = get_logger(__name__)

# Optional pdfplumber for enhanced table extraction
try:
    import pdfplumber

    PDFPLUMBER_AVAILABLE = True
except ImportError:
    pdfplumber = None
    PDFPLUMBER_AVAILABLE = False
    logger.info(
        "pdfplumber_not_available", reason="pdfplumber not installed, table extraction limited"
    )

# Supported MIME types
SUPPORTED_MIME_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "doc",  # Legacy .doc files
}

# Max file size: 50 MB
MAX_FILE_SIZE_BYTES = 50 * 1024 * 1024


@dataclass
class ExtractedTable:
    """A table extracted from a document."""

    page_number: int
    rows: list[list[str]]
    header: list[str] | None = None

    def to_markdown(self) -> str:
        """Convert table to markdown format."""
        if not self.rows:
            return ""

        lines = []
        # Use first row as header if no explicit header
        header = self.header if self.header else self.rows[0]
        data_rows = self.rows if self.header else self.rows[1:]

        # Header row
        lines.append("| " + " | ".join(str(cell) for cell in header) + " |")
        # Separator
        lines.append("| " + " | ".join("---" for _ in header) + " |")
        # Data rows
        for row in data_rows:
            # Pad row if needed
            padded = list(row) + [""] * (len(header) - len(row))
            lines.append("| " + " | ".join(str(cell) for cell in padded[: len(header)]) + " |")

        return "\n".join(lines)


@dataclass
class ExtractedDocument:
    """Result of document text extraction."""

    text: str
    page_count: int
    file_hash: str
    file_size_bytes: int
    mime_type: str
    tables: list[ExtractedTable] = field(default_factory=list)


class DocumentExtractor:
    """Extracts text from PDF and DOCX documents."""

    def extract(
        self,
        content: bytes,
        filename: str,
        mime_type: str,
    ) -> ExtractedDocument:
        """Extract text from a document.

        Args:
            content: Raw file content as bytes
            filename: Original filename
            mime_type: MIME type of the file

        Returns:
            ExtractedDocument with text and metadata

        Raises:
            FileTooLargeError: If file exceeds size limit
            UnsupportedFileTypeError: If file type is not supported
            ValidationError: If document is empty or corrupted
        """
        # Check file size
        file_size = len(content)
        if file_size > MAX_FILE_SIZE_BYTES:
            from dealguard.shared.exceptions import FileTooLargeError

            raise FileTooLargeError(max_size_mb=50)

        # Check MIME type
        if mime_type not in SUPPORTED_MIME_TYPES:
            raise UnsupportedFileTypeError(
                file_type=mime_type,
                supported=list(SUPPORTED_MIME_TYPES.keys()),
            )

        # Calculate file hash
        file_hash = hashlib.sha256(content).hexdigest()

        # Extract based on type
        doc_type = SUPPORTED_MIME_TYPES[mime_type]
        tables: list[ExtractedTable] = []

        if doc_type == "pdf":
            text, page_count, tables = self._extract_pdf(content)
        elif doc_type in ("docx", "doc"):
            text, page_count, tables = self._extract_docx(content)
        else:
            raise UnsupportedFileTypeError(
                file_type=doc_type,
                supported=["pdf", "docx"],
            )

        # Validate extraction
        if not text or len(text.strip()) < 50:
            raise ValidationError("Dokument konnte nicht gelesen werden oder enthält zu wenig Text")

        logger.info(
            "document_extracted",
            filename=filename,
            mime_type=mime_type,
            page_count=page_count,
            text_length=len(text),
            table_count=len(tables),
        )

        return ExtractedDocument(
            text=text,
            page_count=page_count,
            file_hash=file_hash,
            file_size_bytes=file_size,
            mime_type=mime_type,
            tables=tables,
        )

    def _extract_pdf(self, content: bytes) -> tuple[str, int, list[ExtractedTable]]:
        """Extract text and tables from PDF using PyMuPDF and pdfplumber."""
        tables: list[ExtractedTable] = []

        try:
            doc = fitz.open(stream=content, filetype="pdf")
            page_count = len(doc)
            text_parts = []

            for page_num, page in enumerate(doc, 1):
                page_text = page.get_text("text")

                # If no text found, try OCR (for scanned documents)
                if not page_text.strip() and self._should_try_ocr(page):
                    page_text = self._ocr_page(page)

                if page_text:
                    text_parts.append(f"--- Seite {page_num} ---\n{page_text}")

            doc.close()

            # Extract tables with pdfplumber (better table detection)
            if PDFPLUMBER_AVAILABLE:
                tables = self._extract_pdf_tables(content)

            return "\n\n".join(text_parts), page_count, tables

        except Exception as e:
            logger.error("pdf_extraction_failed", error=str(e))
            raise ValidationError(f"PDF konnte nicht gelesen werden: {e}")

    def _extract_pdf_tables(self, content: bytes) -> list[ExtractedTable]:
        """Extract tables from PDF using pdfplumber for better accuracy."""
        tables: list[ExtractedTable] = []

        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_tables = page.extract_tables()

                    for table_data in page_tables:
                        if table_data and len(table_data) > 1:
                            # Clean up table data (replace None with empty string)
                            cleaned_rows = [
                                [str(cell) if cell is not None else "" for cell in row]
                                for row in table_data
                            ]

                            tables.append(
                                ExtractedTable(
                                    page_number=page_num,
                                    rows=cleaned_rows,
                                )
                            )

            logger.info("pdf_tables_extracted", table_count=len(tables))

        except Exception as e:
            logger.warning("pdf_table_extraction_failed", error=str(e))
            # Don't fail completely, just return empty tables

        return tables

    def _extract_docx(self, content: bytes) -> tuple[str, int, list[ExtractedTable]]:
        """Extract text and tables from DOCX."""
        tables: list[ExtractedTable] = []

        try:
            doc = Document(io.BytesIO(content))
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables with structured data
            for table in doc.tables:
                table_rows: list[list[str]] = []

                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    table_rows.append(row_cells)

                    # Also add to text for backward compatibility
                    row_text = " | ".join(cell for cell in row_cells if cell)
                    if row_text:
                        text_parts.append(row_text)

                if table_rows:
                    tables.append(
                        ExtractedTable(
                            page_number=1,  # DOCX doesn't have page numbers easily
                            rows=table_rows,
                        )
                    )

            # Estimate page count (rough approximation)
            text = "\n".join(text_parts)
            estimated_pages = max(1, len(text) // 3000)

            logger.info("docx_tables_extracted", table_count=len(tables))

            return text, estimated_pages, tables

        except Exception as e:
            logger.error("docx_extraction_failed", error=str(e))
            raise ValidationError(f"DOCX konnte nicht gelesen werden: {e}")

    def _should_try_ocr(self, page: fitz.Page) -> bool:
        """Check if page appears to be a scanned image."""
        # Get images on page
        image_list = page.get_images()
        # If page has large images but no text, it's likely scanned
        return len(image_list) > 0

    def _ocr_page(self, page: fitz.Page) -> str:
        """Perform OCR on a page using Tesseract.

        This is optional and only runs if Tesseract is installed.
        """
        try:
            import pytesseract
            from PIL import Image

            # Render page to image
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)

            # OCR with German language
            text = pytesseract.image_to_string(img, lang="deu")
            return cast(str, text)

        except ImportError:
            logger.warning("ocr_not_available", reason="pytesseract not installed")
            return ""
        except Exception as e:
            logger.warning("ocr_failed", error=str(e))
            return ""
